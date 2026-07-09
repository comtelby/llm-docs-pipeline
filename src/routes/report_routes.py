import io
import markdown
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse, PlainTextResponse, HTMLResponse

import aiofiles

from src.config import OUTPUT_DIR
from src.database import list_audit_history
from src.report import generate_report

from src.state import get_last_prompt

router = APIRouter()


@router.post("/generate-report")
async def create_report():
    prompt = get_last_prompt()
    if not prompt:
        raise HTTPException(400, "Сначала сформулируйте задачу в чате")
    result = await generate_report(prompt)
    return result


@router.get("/reports")
async def list_reports():
    reports = []
    if OUTPUT_DIR.exists():
        for f in OUTPUT_DIR.iterdir():
            if f.is_file():
                s = f.stat()
                from datetime import datetime
                reports.append({
                    "name": f.name,
                    "size_bytes": s.st_size,
                    "modified": datetime.fromtimestamp(s.st_mtime).isoformat()
                })
    return {"reports": sorted(reports, key=lambda x: x["modified"], reverse=True)}


@router.get("/download/{filename}")
async def download_report(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    return FileResponse(str(path), filename=filename)


@router.get("/view/{filename}")
async def view_report(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    async with aiofiles.open(path, "r", encoding="utf-8") as f:
        content = await f.read()
    return {"filename": filename, "content": content}


@router.get("/export/{filename}")
async def export_report(filename: str, format: str = "md"):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404, "Файл не найден")
    async with aiofiles.open(path, "r", encoding="utf-8") as f:
        content = await f.read()

    if format == "md":
        return FileResponse(str(path), filename=filename)
    elif format == "txt":
        return PlainTextResponse(
            content,
            headers={"Content-Disposition": f"attachment; filename={filename.replace('.md', '.txt')}"}
        )
    elif format == "html":
        html = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        return HTMLResponse(
            f"<html><head><meta charset='UTF-8'>"
            f"<style>body{{font-family:Arial;max-width:900px;margin:40px auto;padding:20px}}"
            f"table{{border-collapse:collapse;width:100%}}"
            f"td,th{{border:1px solid #ddd;padding:8px}}</style>"
            f"</head><body>{html}</body></html>"
        )
    elif format == "docx":
        return await _export_docx(content, filename)
    elif format == "pdf":
        return await _export_pdf(content, filename)
    else:
        raise HTTPException(400, f"Формат {format} не поддерживается")


def _md_to_docx(content: str, buf: io.BytesIO):
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        if line.startswith('# ') or line.startswith('## ') or line.startswith('### '):
            level = len(line) - len(line.lstrip('#'))
            heading_text = line.lstrip('# ').strip()
            if level <= 1:
                doc.add_heading(heading_text, level=1)
            elif level == 2:
                doc.add_heading(heading_text, level=2)
            else:
                doc.add_heading(heading_text, level=3)

        elif line.startswith('|') and line.endswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            i -= 1
            if len(rows) >= 2:
                table = doc.add_table(rows=len(rows) - 1, cols=len(rows[0].split('|')) - 2)
                table.style = 'Light Grid Accent 1'
                for r_idx, row_text in enumerate(rows):
                    if r_idx == 1 and set(row_text.replace('|', '').replace('-', '').strip()) == set():
                        continue
                    cells = [c.strip() for c in row_text.split('|')[1:-1]]
                    for c_idx, cell_text in enumerate(cells):
                        if r_idx == 0:
                            cell = table.rows[0].cells[c_idx]
                        else:
                            cell = table.rows[r_idx - 1].cells[c_idx]
                        cell.text = cell_text

        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            run = p.add_run('\n'.join(code_lines))
            run.font.size = Pt(9)
            run.font.name = 'Consolas'

        elif line.startswith('---'):
            doc.add_paragraph('─' * 40)

        elif line.startswith('*') and line.endswith('*'):
            doc.add_paragraph(line.strip('*')).italic = True

        else:
            p = doc.add_paragraph(line)
            if line.startswith('**') and line.endswith('**'):
                p.clear()
                run = p.add_run(line.strip('*'))
                run.bold = True

        i += 1

    doc.save(buf)


async def _export_docx(content: str, filename: str) -> FileResponse:
    buf = io.BytesIO()
    _md_to_docx(content, buf)
    buf.seek(0)
    docx_filename = filename.replace('.md', '.docx')
    tmp_path = OUTPUT_DIR / docx_filename
    with open(tmp_path, 'wb') as f:
        f.write(buf.read())
    return FileResponse(str(tmp_path), filename=docx_filename)


def _find_dejavu_fonts():
    import os
    paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
        '/usr/local/share/fonts/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/dejavu/DejaVuSans.ttf',
    ]
    regular = None
    bold = None
    for p in paths:
        if os.path.exists(p):
            if 'Bold' in p:
                bold = p
            else:
                regular = p
    return regular, bold


async def _export_pdf(content: str, filename: str) -> FileResponse:
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(500, "fpdf2 не установлен")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    font_regular, font_bold = _find_dejavu_fonts()
    if font_regular:
        pdf.add_font('DejaVu', '', font_regular, uni=True)
        if font_bold:
            pdf.add_font('DejaVu', 'B', font_bold, uni=True)
        has_cyrillic = True
    else:
        has_cyrillic = False

    lines = content.split('\n')
    in_code = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            pdf.ln(3)
            continue

        if stripped.startswith('```'):
            in_code = not in_code
            continue

        if in_code:
            if has_cyrillic:
                pdf.set_font('DejaVu', '', 8)
            else:
                pdf.set_font('Courier', '', 8)
            pdf.multi_cell(0, 4, stripped)
            continue

        if has_cyrillic:
            if stripped.startswith('# ') or stripped.startswith('## ') or stripped.startswith('### '):
                level = len(stripped) - len(stripped.lstrip('#'))
                text = stripped.lstrip('# ').strip()
                pdf.set_font('DejaVu', 'B', max(12, 16 - level * 2))
                pdf.multi_cell(0, 8, text)
                pdf.ln(2)
            elif stripped.startswith('|') and stripped.endswith('|'):
                pdf.set_font('DejaVu', '', 8)
                cells = [c.strip() for c in stripped.split('|')[1:-1]]
                col_w = 180 / max(len(cells), 1)
                for c in cells:
                    pdf.cell(col_w, 6, c, border=1)
                pdf.ln()
            elif stripped.startswith('---'):
                pdf.ln(3)
            else:
                pdf.set_font('DejaVu', '', 10)
                pdf.multi_cell(0, 5, stripped)
        else:
            safe = stripped.encode('ascii', errors='replace').decode('ascii')
            pdf.set_font('Helvetica', '', 10)
            pdf.multi_cell(0, 5, safe)

    pdf_filename = filename.replace('.md', '.pdf')
    tmp_path = OUTPUT_DIR / pdf_filename
    pdf.output(str(tmp_path))
    return FileResponse(str(tmp_path), filename=pdf_filename)


@router.get("/audit-history")
async def get_audit_history():
    return {"history": list_audit_history()}
