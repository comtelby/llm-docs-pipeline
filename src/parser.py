import re
import subprocess
import logging
from pathlib import Path

from src.eol import lookup_eol
from src.models import DeviceInfo

logger = logging.getLogger(__name__)

MODEL_MAP = {
    "hp5412": "HP 5412R 92GT PoE+",
    "hp-5412": "HP 5412R 92GT PoE+",
    "hp5510": "HPE 5510 24G SFP 4SFP+ HI",
    "5510_m1": "HPE 5510 24G SFP 4SFP+ HI",
    "s6730": "Huawei S6730-H24X6C",
    "h6121": "Huawei AR6121E",
    "c4331": "Cisco ISR4331-K9",
    "asa5516": "Cisco ASA 5516-X",
    "2530-48": "HP Aruba 2530 48 PoE+",
    "2530": "HP Aruba 2530 24 PoE+",
    "hpe5120": "HPE 5120 16G",
    "5120-02": "HPE 5120 16G",
    "hp3500": "HP ProCurve 3500-24-PoE",
    "hp2620": "HP 2620-24 PoE+",
    "5731": "Huawei 5731-H48P4XC",
    "ac6508": "Huawei AC6508",
    "hwlc01": "Huawei AC6508",
    "by_core": "Cisco ISR4331/K9",
}


def parse_device_config(text: str, filename: str) -> DeviceInfo:
    device = DeviceInfo(file=filename, full_length=len(text))

    hostname_match = re.search(r'(?:hostname|sysname)\s+["\']?([^"\'\n\r]+)["\']?', text, re.IGNORECASE)
    if hostname_match:
        device.hostname = hostname_match.group(1).strip()

    hostname_lower = device.hostname.lower()
    for key, model_name in MODEL_MAP.items():
        if key in hostname_lower:
            device.model = model_name
            break

    device.eol_info = lookup_eol(device.hostname, device.model)

    mgmt_match = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', text, re.IGNORECASE)
    if mgmt_match:
        device.ip_mgmt = mgmt_match.group(1)

    device.ip_addresses = re.findall(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', text, re.IGNORECASE)[:50]

    interface_blocks = re.findall(
        r'^interface\s+(.+?)$(.*?)(?=^interface\s|\Z)',
        text, re.MULTILINE | re.DOTALL | re.IGNORECASE
    )
    for iface_name, iface_config in interface_blocks[:30]:
        iface_ip = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', iface_config, re.IGNORECASE)
        device.interfaces.append({
            "name": iface_name.strip(),
            "ip": f"{iface_ip.group(1)}/{iface_ip.group(2)}" if iface_ip else "Нет IP"
        })

    vlan_blocks = re.findall(r'^vlan\s+(\d+).*$(.*?)(?=^vlan|\Z)', text, re.MULTILINE | re.DOTALL | re.IGNORECASE)
    for vlan_id, vlan_config in vlan_blocks[:30]:
        vlan_ip = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', vlan_config, re.IGNORECASE)
        device.vlans.append({
            "id": vlan_id,
            "ip": f"{vlan_ip.group(1)}/{vlan_ip.group(2)}" if vlan_ip else "L2"
        })

    ospf_process = re.search(r"ospf\s+(\d+)", text, re.IGNORECASE)
    ospf_router_id = re.search(r"router-id\s+([\d.]+)", text, re.IGNORECASE)
    ospf_area = re.search(r"area\s+([\d.]+)", text, re.IGNORECASE)
    if ospf_process or ospf_router_id:
        device.ospf = {
            "process_id": ospf_process.group(1) if ospf_process else "?",
            "router_id": ospf_router_id.group(1) if ospf_router_id else "?",
            "area": ospf_area.group(1) if ospf_area else "?"
        }

    vrrp_blocks = re.findall(r'vrrp\s+(\d+)\s+.*?virtual-ip\s+([\d.]+)', text, re.IGNORECASE)
    device.vrrp = [{"group": g, "vip": vip} for g, vip in vrrp_blocks[:10]]

    stp_region = re.search(r'region-name\s+(\S+)', text, re.IGNORECASE)
    if stp_region:
        device.stp_region = stp_region.group(1)
    device.stp_instances = re.findall(r'instance\s+(\d+)\s+vlan\s+(.+)', text, re.IGNORECASE)[:10]

    device.tacacs_servers = re.findall(r'tacacs-server\s+host\s+([\d.]+)', text, re.IGNORECASE)[:5]
    device.syslog_servers = re.findall(r'syslog.*?([\d.]+)', text, re.IGNORECASE)[:5]
    device.ntp_servers = re.findall(r'ntp\s+(?:server|source)\s+([\d.]+)', text, re.IGNORECASE)[:5]

    device.dhcp_snooping = bool(re.search(r'dhcp-snooping', text, re.IGNORECASE))
    device.port_security = bool(re.search(r'port-security', text, re.IGNORECASE))
    device.bpdu_protection = bool(re.search(r'bpdu-protection', text, re.IGNORECASE))

    device.acl = re.findall(r'^(?:access-list|acl)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:20]
    device.aaa = re.findall(r'^(?:aaa|radius|tacacs)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:20]
    device.snmp = re.findall(r'^(?:snmp-server|snmp-agent)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:10]
    device.routes = re.findall(r'^ip\s+route\s+(.+)$', text, re.MULTILINE | re.IGNORECASE)[:30]

    return device


def extract_text_from_image(image_path: Path) -> str:
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='rus+eng')
        if text.strip():
            logger.info(f"OCR OK: {image_path.name} ({len(text)} chars)")
            return text.strip()
        logger.warning(f"OCR пустой результат: {image_path.name}")
        return ""
    except ImportError as e:
        logger.error(f"OCR: нет библиотеки: {e}")
        return ""
    except Exception as e:
        logger.error(f"OCR ошибка {image_path.name}: {e}")
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(docx_path))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
        result = '\n'.join(paragraphs)
        logger.info(f"DOCX: {docx_path.name} ({len(result)} chars)")
        return result
    except ImportError:
        logger.error("DOCX: python-docx не установлен")
        return ""
    except Exception as e:
        logger.error(f"DOCX ошибка {docx_path.name}: {e}")
        return ""


def extract_text_from_doc(doc_path: Path) -> str:
    """Извлечение текста из старого .doc (Word 97-2003)."""
    # Попытка 1: antiword
    try:
        result = subprocess.run(
            ["antiword", "-m", "UTF-8", str(doc_path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            logger.info(f"DOC(antiword): {doc_path.name} ({len(result.stdout)} chars)")
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.warning(f"DOC antiword error {doc_path.name}: {e}")

    # Попытка 2: catdoc
    try:
        result = subprocess.run(
            ["catdoc", str(doc_path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout.strip():
            logger.info(f"DOC(catdoc): {doc_path.name} ({len(result.stdout)} chars)")
            return result.stdout.strip()
    except FileNotFoundError:
        pass
    except Exception as e:
        logger.warning(f"DOC catdoc error {doc_path.name}: {e}")

    # Попытка 3: raw text через olefile
    try:
        import olefile
        ole = olefile.OleFileIO(str(doc_path))
        if ole.exists('WordDocument'):
            stream = ole.openstream('WordDocument')
            raw = stream.read()
            text = raw.decode('utf-8', errors='ignore')
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
            text = re.sub(r'\s+', ' ', text).strip()
            if len(text) > 100:
                logger.info(f"DOC(olefile): {doc_path.name} ({len(text)} chars)")
                return text[:50000]
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"DOC olefile error {doc_path.name}: {e}")

    logger.warning(f"DOC: не удалось извлечь текст из {doc_path.name}")
    return ""


def extract_text_from_rtf(rtf_path: Path) -> str:
    try:
        # Читаем как latin-1, чтобы сохранить \'xx escape-последовательности для striprtf
        raw_bytes = rtf_path.read_bytes()
        raw_text = raw_bytes.decode('latin-1')

        # Попытка 1: striprtf (правильный парсинг RTF)
        try:
            from striprtf.striprtf import rtf_to_text
            result = rtf_to_text(raw_text)
            if len(result.strip()) > 100:
                logger.info(f"RTF(striprtf): {rtf_path.name} ({len(result)} chars)")
                return result.strip()
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"RTF striprtf error {rtf_path.name}: {e}")

        # Попытка 2: ручная очистка (fallback)
        content = raw_text
        content = re.sub(r'\\[a-z]+\d*', '', content)
        content = re.sub(r'[\\{};]', '', content)
        content = content.replace('\\par', '\n')
        content = re.sub(r"\'[0-9a-fA-F]{2}", lambda m: chr(int(m.group(0)[1:], 16)), content)
        lines = [line.strip() for line in content.split('\n') if line.strip() and len(line.strip()) > 3]
        result = '\n'.join(lines)
        logger.info(f"RTF(fallback): {rtf_path.name} ({len(result)} chars)")
        return result
    except Exception as e:
        logger.error(f"RTF ошибка {rtf_path.name}: {e}")
        return ""


def extract_text_from_xlsx(xlsx_path: Path) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_str = ' | '.join(str(c) if c is not None else '' for c in row)
                if row_str.strip():
                    rows_text.append(row_str)
            if rows_text:
                parts.append(f"=== {sheet_name} ===\n" + '\n'.join(rows_text))
        result = '\n\n'.join(parts)
        logger.info(f"XLSX: {xlsx_path.name} ({len(result)} chars)")
        return result
    except ImportError:
        logger.error("XLSX: openpyxl не установлен")
        return ""
    except Exception as e:
        logger.error(f"XLSX ошибка {xlsx_path.name}: {e}")
        return ""


def extract_text_from_xls(xls_path: Path) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(str(xls_path))
        parts = []
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            rows_text = []
            for row_idx in range(ws.nrows):
                row_str = ' | '.join(str(ws.cell_value(row_idx, c)) for c in range(ws.ncols))
                if row_str.strip():
                    rows_text.append(row_str)
            if rows_text:
                parts.append(f"=== {sheet_name} ===\n" + '\n'.join(rows_text))
        result = '\n\n'.join(parts)
        logger.info(f"XLS: {xls_path.name} ({len(result)} chars)")
        return result
    except ImportError:
        logger.error("XLS: xlrd не установлен, пробую pandas")
        try:
            import pandas as pd
            dfs = pd.read_excel(xls_path, sheet_name=None)
            parts = []
            for sheet_name, df in dfs.items():
                rows_text = df.astype(str).to_csv(sep=' | ', index=False)
                if rows_text.strip():
                    parts.append(f"=== {sheet_name} ===\n{rows_text}")
            result = '\n\n'.join(parts)
            logger.info(f"XLS(pandas): {xls_path.name} ({len(result)} chars)")
            return result
        except ImportError:
            logger.error("XLS: ни xlrd, ни pandas недоступны")
            return ""
        except Exception as e:
            logger.error(f"XLS pandas ошибка {xls_path.name}: {e}")
            return ""
    except Exception as e:
        logger.error(f"XLS ошибка {xls_path.name}: {e}")
        return ""


def extract_text_from_file(file_path: Path) -> str:
    """Универсальное извлечение текста из файла любого поддерживаемого формата."""
    suffix = file_path.suffix.lower()
    try:
        if suffix == '.docx':
            return extract_text_from_docx(file_path)
        elif suffix == '.doc':
            return extract_text_from_doc(file_path)
        elif suffix == '.xlsx':
            return extract_text_from_xlsx(file_path)
        elif suffix == '.xls':
            return extract_text_from_xls(file_path)
        elif suffix == '.rtf':
            return extract_text_from_rtf(file_path)
        elif suffix in {'.txt', '.md', '.cfg', '.conf'}:
            return read_text_file(file_path)
        elif suffix == '.csv':
            return read_text_file(file_path)
        else:
            logger.debug(f"extract_text_from_file: неподдерживаемый формат {suffix} для {file_path.name}")
            return ""
    except Exception as e:
        logger.warning(f"extract_text_from_file ошибка {file_path.name}: {e}")
        return ""


def read_text_file(file_path: Path) -> str:
    for encoding in ['utf-8', 'cp1251', 'latin-1', 'ascii']:
        try:
            return file_path.read_text(encoding=encoding)
        except (UnicodeDecodeError, FileNotFoundError):
            continue
    return ""
