import os
import shutil
import uuid
import json
import requests
import markdown
import re
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, Request, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, PlainTextResponse
from jinja2 import Environment, FileSystemLoader, select_autoescape
import aiofiles

app = FastAPI()

# --- Конфигурация путей ---
BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
UPLOAD_DIR = BASE_DIR / "input" / "uploads"
CONFIGS_DIR = BASE_DIR / "input" / "configs"
SCREENSHOTS_DIR = BASE_DIR / "input" / "screenshots"
SAMPLES_DIR = BASE_DIR / "input" / "samples"
OUTPUT_DIR = BASE_DIR / "output"

for directory in [UPLOAD_DIR, CONFIGS_DIR, SCREENSHOTS_DIR, SAMPLES_DIR, OUTPUT_DIR]:
    directory.mkdir(parents=True, exist_ok=True)

env = Environment(
    loader=FileSystemLoader(str(TEMPLATES_DIR)),
    autoescape=select_autoescape(["html", "xml"]),
    cache_size=0
)

OLLAMA_API_URL = os.getenv("OLLAMA_API_URL", "http://ollama:11434/api")
last_user_prompt = ""


# ===== БАЗА ДАННЫХ ОБОРУДОВАНИЯ (EOL/EOSL) =====
EOL_DATABASE = {
    "avaya ip office 500v2": {"eol": "31.12.2020", "status": "EOSL", "note": "Поддержка прекращена"},
    "hp 2620-24": {"eol": "2020", "status": "End-of-Sale", "note": "Обновлений нет"},
    "hp procurve 3500-24": {"eol": "31.10.2019", "status": "EOSL", "note": "Производство прекращено"},
    "hpe 5120 16g": {"eol": "2019", "status": "End-of-Sale", "note": "Заменён на серию 5130"},
    "hpe 5510": {"eol": "2021", "status": "End-of-Sale", "note": "Заменён на серию 5520"},
    "huawei 5731": {"eol": None, "status": "Актуальное", "note": "Поддержка до 2029"},
    "huawei s6730": {"eol": None, "status": "Актуальное", "note": "Поддержка до 2031+"},
    "huawei ac6508": {"eol": None, "status": "Актуальное", "note": "Активная продажа"},
    "huawei ar6121e": {"eol": None, "status": "Актуальное", "note": "Поддержка до 2030+"},
    "huawei oceanstor dorado 2000": {"eol": None, "status": "Актуальное", "note": "Активная продажа"},
    "ibm system x3650 m5": {"eol": "2019", "status": "End-of-Sale", "note": "Поддержка прекращена (Lenovo)"},
    "ibm system x3630 m4": {"eol": "2019", "status": "EOSL", "note": "Полное прекращение поддержки"},
    "synology rs814+": {"eol": "01.10.2024", "status": "EOSL", "note": "DSM 6.2, обновлений нет"},
    "synology rs815+": {"eol": "01.10.2024", "status": "EOSL", "note": "DSM 6.2, обновлений нет"},
    "synology rx415": {"eol": "01.10.2024", "status": "EOSL", "note": "Не поддерживает DSM 7.x"},
    "synology rx1217": {"eol": "2023", "status": "End-of-Sale", "note": "Снят с производства"},
    "synology rs3618xs": {"eol": "2022", "status": "End-of-Sale", "note": "Снят с продаж"},
    "synology sa3200d": {"eol": None, "status": "Актуальное", "note": "Активная продажа"},
    "cisco asa 5516": {"eol": "2022", "status": "End-of-Sale", "note": "Последний патч: август 2024"},
    "cisco isr4331": {"eol": "2022", "status": "End-of-Sale", "note": "SW поддержка до 2026"},
    "cisco air-ct-3504": {"eol": "2021", "status": "End-of-Sale", "note": "Заменён на 3504"},
    "hp 5412r": {"eol": "2020", "status": "End-of-Sale", "note": "Поддержка ограничена"},
    "hp aruba 2530": {"eol": "2021", "status": "End-of-Sale", "note": "Заменён на 2540"},
    "lenovo b6505": {"eol": "2020", "status": "End-of-Sale", "note": "Снят с производства"},
    "lenovo ds2200": {"eol": "2023", "status": "EOSL", "note": "Полное прекращение поддержки"},
    "dell r230": {"eol": "31.03.2023", "status": "EOSL", "note": "iDRAC 8, обновлений нет"},
    "ibm ts3200": {"eol": "31.12.2023", "status": "EOSL", "note": "Поддержка прекращена"},
    "oring rgs-7168": {"eol": "2023", "status": "EOSL", "note": "Производство прекращено"},
    "apc smart-ups 2200": {"eol": "2022", "status": "EOSL", "note": "Снят с производства"},
    "apc smart-ups srt 5000": {"eol": None, "status": "Актуальное", "note": "В активной продаже"},
    "depo cs-3400": {"eol": "Неизвестно", "status": "End-of-Sale", "note": "Устаревшая модель"},
}


def lookup_eol(hostname, model=""):
    """Поиск EOL-статуса по hostname или модели"""
    search_terms = []
    if hostname:
        search_terms.append(hostname.lower())
    if model:
        search_terms.append(model.lower())
    
    for term in search_terms:
        for key, value in EOL_DATABASE.items():
            if key in term or term in key:
                return value
    
    return {"eol": "Неизвестно", "status": "Требуется проверка", "note": "Нет в базе данных"}


def parse_device_config(text, filename):
    """Глубокий парсинг конфигурации устройства"""
    device = {
        "file": filename,
        "hostname": "",
        "model": "",
        "ip_mgmt": "",
        "ip_addresses": [],
        "interfaces": [],
        "vlans": [],
        "routes": [],
        "aaa": [],
        "acl": [],
        "ntp": [],
        "snmp": [],
        "spanning_tree": [],
        "ospf": {},
        "vrrp": [],
        "stp_region": "",
        "stp_instances": [],
        "tacacs_servers": [],
        "syslog_servers": [],
        "dhcp_snooping": False,
        "port_security": False,
        "bpdu_protection": False,
        "ntp_servers": [],
        "eol_info": {},
        "full_length": len(text)
    }
    
    # Hostname
    hostname_match = re.search(r'(?:hostname|sysname)\s+["\']?([^"\'\n\r]+)["\']?', text, re.IGNORECASE)
    if hostname_match:
        device["hostname"] = hostname_match.group(1).strip()
    
    # Модель из hostname
    hostname_lower = device["hostname"].lower()
    if "hp5412" in hostname_lower:
        device["model"] = "HP 5412R 92GT PoE+"
    elif "hp5510" in hostname_lower or "5510_m1" in hostname_lower:
        device["model"] = "HPE 5510 24G SFP 4SFP+ HI"
    elif "s6730" in hostname_lower:
        device["model"] = "Huawei S6730-H24X6C"
    elif "h6121" in hostname_lower:
        device["model"] = "Huawei AR6121E"
    elif "c4331" in hostname_lower:
        device["model"] = "Cisco ISR4331-K9"
    elif "asa5516" in hostname_lower:
        device["model"] = "Cisco ASA 5516-X"
    elif "2530-48" in hostname_lower:
        device["model"] = "HP Aruba 2530 48 PoE+"
    elif "2530" in hostname_lower:
        device["model"] = "HP Aruba 2530 24 PoE+"
    elif "hpe5120" in hostname_lower or "5120-02" in hostname_lower:
        device["model"] = "HPE 5120 16G"
    elif "hp3500" in hostname_lower:
        device["model"] = "HP ProCurve 3500-24-PoE"
    elif "hp2620" in hostname_lower:
        device["model"] = "HP 2620-24 PoE+"
    elif "5731" in hostname_lower:
        device["model"] = "Huawei 5731-H48P4XC"
    elif "ac6508" in hostname_lower or "hwlc01" in hostname_lower:
        device["model"] = "Huawei AC6508"
    
    # EOL
    device["eol_info"] = lookup_eol(device["hostname"], device["model"])
    
    # IP management
    mgmt_match = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', text, re.IGNORECASE)
    if mgmt_match:
        device["ip_mgmt"] = mgmt_match.group(1)
    
    # IP-адреса
    device["ip_addresses"] = re.findall(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', text, re.IGNORECASE)[:50]
    
    # Интерфейсы
    interface_blocks = re.findall(
        r'^interface\s+(.+?)$(.*?)(?=^interface\s|\Z)', 
        text, re.MULTILINE | re.DOTALL | re.IGNORECASE
    )
    for iface_name, iface_config in interface_blocks[:30]:
        iface_ip = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', iface_config, re.IGNORECASE)
        device["interfaces"].append({
            "name": iface_name.strip(),
            "ip": f"{iface_ip.group(1)}/{iface_ip.group(2)}" if iface_ip else "Нет IP"
        })
    
    # VLAN
    vlan_blocks = re.findall(r'^vlan\s+(\d+).*$(.*?)(?=^vlan|\Z)', text, re.MULTILINE | re.DOTALL | re.IGNORECASE)
    for vlan_id, vlan_config in vlan_blocks[:30]:
        vlan_ip = re.search(r'ip\s+address\s+([\d.]+)\s+([\d.]+)', vlan_config, re.IGNORECASE)
        device["vlans"].append({
            "id": vlan_id,
            "ip": f"{vlan_ip.group(1)}/{vlan_ip.group(2)}" if vlan_ip else "L2"
        })
    
    # OSPF (ИСПРАВЛЕНО)
    ospf_process = re.search(r"ospf\s+(\d+)", text, re.IGNORECASE)
    ospf_router_id = re.search(r"router-id\s+([\d.]+)", text, re.IGNORECASE)
    ospf_area = re.search(r"area\s+([\d.]+)", text, re.IGNORECASE)
    if ospf_process or ospf_router_id:
        device["ospf"] = {
            "process_id": ospf_process.group(1) if ospf_process else "?",
            "router_id": ospf_router_id.group(1) if ospf_router_id else "?",
            "area": ospf_area.group(1) if ospf_area else "?"
        }
    
    # VRRP
    vrrp_blocks = re.findall(r'vrrp\s+(\d+)\s+.*?virtual-ip\s+([\d.]+)', text, re.IGNORECASE)
    device["vrrp"] = [{"group": g, "vip": vip} for g, vip in vrrp_blocks[:10]]
    
    # STP
    stp_region = re.search(r'region-name\s+(\S+)', text, re.IGNORECASE)
    if stp_region:
        device["stp_region"] = stp_region.group(1)
    device["stp_instances"] = re.findall(r'instance\s+(\d+)\s+vlan\s+(.+)', text, re.IGNORECASE)[:10]
    
    # TACACS
    device["tacacs_servers"] = re.findall(r'tacacs-server\s+host\s+([\d.]+)', text, re.IGNORECASE)[:5]
    
    # Syslog
    device["syslog_servers"] = re.findall(r'syslog.*?([\d.]+)', text, re.IGNORECASE)[:5]
    
    # NTP
    device["ntp_servers"] = re.findall(r'ntp\s+(?:server|source)\s+([\d.]+)', text, re.IGNORECASE)[:5]
    
    # Безопасность
    device["dhcp_snooping"] = bool(re.search(r'dhcp-snooping', text, re.IGNORECASE))
    device["port_security"] = bool(re.search(r'port-security', text, re.IGNORECASE))
    device["bpdu_protection"] = bool(re.search(r'bpdu-protection', text, re.IGNORECASE))
    
    # ACL
    device["acl"] = re.findall(r'^(?:access-list|acl)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:20]
    
    # AAA
    device["aaa"] = re.findall(r'^(?:aaa|radius|tacacs)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:20]
    
    # SNMP
    device["snmp"] = re.findall(r'^(?:snmp-server|snmp-agent)\s+.+$', text, re.MULTILINE | re.IGNORECASE)[:10]
    
    # Маршруты
    device["routes"] = re.findall(r'^ip\s+route\s+(.+)$', text, re.MULTILINE | re.IGNORECASE)[:30]
    
    return device


def extract_text_from_image(image_path):
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='rus+eng')
        return text.strip() if text.strip() else ""
    except:
        return ""


def extract_text_from_docx(docx_path):
    try:
        from docx import Document
        doc = Document(str(docx_path))
        paragraphs = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                if row_text.strip():
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs)
    except:
        return ""


def extract_text_from_rtf(rtf_path):
    try:
        with open(rtf_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        content = re.sub(r'\\[a-z]+\d*', '', content)
        content = re.sub(r'[\\{};]', '', content)
        lines = [line.strip() for line in content.split('\n') if line.strip() and len(line.strip()) > 3]
        return '\n'.join(lines)
    except:
        return ""


def read_text_file(file_path):
    for encoding in ['utf-8', 'cp1251', 'latin-1', 'ascii']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except:
            continue
    return ""


# ===== ЭНДПОИНТЫ =====

@app.get("/health")
async def health():
    return {"status": "ok", "message": "Server is running"}


@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    try:
        template = env.get_template("index.html")
        return template.render({"request": request, "title": "iqData Bot - Аудит ИТ-инфраструктуры"})
    except:
        return HTMLResponse("<h1>Ошибка загрузки шаблона</h1>", status_code=500)


@app.get("/models")
async def list_models():
    try:
        resp = requests.get(f"{OLLAMA_API_URL}/tags", timeout=10)
        return {"models": [m["name"] for m in resp.json().get("models", [])]}
    except:
        return {"error": "Ollama недоступна"}, 503


@app.get("/files")
async def list_files():
    files_info = {"inventory": [], "configs": [], "screenshots": [], "samples": [], "reports": []}
    for category, directory in [
        ("inventory", UPLOAD_DIR), ("configs", CONFIGS_DIR),
        ("screenshots", SCREENSHOTS_DIR), ("samples", SAMPLES_DIR),
        ("reports", OUTPUT_DIR)
    ]:
        if directory.exists():
            for f in directory.iterdir():
                if f.is_file():
                    s = f.stat()
                    files_info[category].append({
                        "name": f.name, "size_bytes": s.st_size,
                        "modified": datetime.fromtimestamp(s.st_mtime).isoformat()
                    })
    return files_info


@app.post("/upload/{category}")
async def upload_file_category(category: str, file: UploadFile = File(...)):
    dir_map = {"inventory": UPLOAD_DIR, "configs": CONFIGS_DIR, "screenshots": SCREENSHOTS_DIR, "samples": SAMPLES_DIR}
    if category not in dir_map:
        raise HTTPException(400, "Неизвестная категория")
    path = dir_map[category] / f"{uuid.uuid4().hex}_{file.filename}"
    with open(path, "wb") as b:
        shutil.copyfileobj(file.file, b)
    return {"status": "success", "filename": file.filename}


@app.delete("/files/{category}/{filename}")
async def delete_file(category: str, filename: str):
    dir_map = {"inventory": UPLOAD_DIR, "configs": CONFIGS_DIR, "screenshots": SCREENSHOTS_DIR, "samples": SAMPLES_DIR, "reports": OUTPUT_DIR}
    if category not in dir_map:
        raise HTTPException(400, "Неизвестная категория")
    path = dir_map[category] / filename
    if path.exists():
        os.remove(path)
    return {"status": "deleted"}


@app.post("/clear-temp")
async def clear_temp():
    c = {"configs": 0, "screenshots": 0}
    for d, k in [(CONFIGS_DIR, "configs"), (SCREENSHOTS_DIR, "screenshots")]:
        if d.exists():
            for f in d.iterdir():
                if f.is_file():
                    os.remove(f)
                    c[k] += 1
    return {"status": "success", "cleared": c}


@app.post("/chat")
async def chat(prompt: str = Form(...), model: str = Form(default="qwen2.5-coder:7b")):
    global last_user_prompt
    last_user_prompt = prompt.strip()
    try:
        r = requests.post(f"{OLLAMA_API_URL}/generate", json={"model": model, "prompt": prompt, "stream": False}, timeout=120)
        return {"response": r.json().get("response", ""), "prompt_saved": True}
    except:
        return {"error": "Ошибка Ollama"}, 500


# ===== ГЛАВНАЯ ФУНКЦИЯ ГЕНЕРАЦИИ ОТЧЁТА =====

@app.post("/generate-report")
async def generate_report():
    global last_user_prompt
    if not last_user_prompt:
        raise HTTPException(400, "Сначала сформулируйте задачу в чате")
    
    print("=" * 60)
    print("STARTING DEEP REPORT GENERATION")
    
    # === ЭТАП 1: ПАРСИНГ ВСЕХ ДАННЫХ ===
    devices = []
    screenshots_data = []
    template_text = ""
    inventory_text = ""
    
    # Конфиги
    if CONFIGS_DIR.exists():
        for f in CONFIGS_DIR.iterdir():
            if f.is_file() and f.suffix.lower() in {".txt", ".cfg", ".conf"}:
                text = read_text_file(f)
                if text and len(text) > 100:
                    device = parse_device_config(text, f.name)
                    if device["hostname"]:
                        devices.append(device)
                        print(f"  ✓ {device['hostname']} ({device['model']}) - EOL: {device['eol_info']['status']}")
    
    # Скриншоты
    if SCREENSHOTS_DIR.exists():
        for f in SCREENSHOTS_DIR.iterdir():
            if f.is_file() and f.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                ocr = extract_text_from_image(str(f))
                if ocr and len(ocr) > 20:
                    screenshots_data.append({"file": f.name, "text": ocr[:2000]})
    
    # Шаблон
    if SAMPLES_DIR.exists():
        for f in SAMPLES_DIR.iterdir():
            if f.is_file() and f.suffix.lower() == '.docx':
                template_text = extract_text_from_docx(f)[:10000]
    
    # Inventory
    if UPLOAD_DIR.exists():
        for f in UPLOAD_DIR.iterdir():
            if f.is_file() and f.suffix.lower() in {".rtf", ".txt", ".md"}:
                if f.suffix.lower() == '.rtf':
                    inventory_text = extract_text_from_rtf(f)[:10000]
                else:
                    inventory_text = read_text_file(f)[:10000]
    
    # === ЭТАП 2: АНАЛИТИКА (без модели) ===
    
    # Группировка устройств
    core_devices = [d for d in devices if any(k in d["hostname"].lower() for k in ["core", "5510", "s6730"])]
    dist_devices = [d for d in devices if any(k in d["hostname"].lower() for k in ["hp5412", "hp3500"])]
    access_devices = [d for d in devices if any(k in d["hostname"].lower() for k in ["2530", "2620", "5120", "5731"])]
    edge_devices = [d for d in devices if any(k in d["hostname"].lower() for k in ["c4331", "h6121", "asa"])]
    
    # EOL анализ
    eol_critical = [d for d in devices if d["eol_info"]["status"] == "EOSL"]
    eol_warning = [d for d in devices if d["eol_info"]["status"] == "End-of-Sale"]
    eol_ok = [d for d in devices if d["eol_info"]["status"] == "Актуальное"]
    
    # Устройства без AAA/NTP/ACL
    no_aaa = [d for d in devices if not d["aaa"]]
    no_ntp = [d for d in devices if not d["ntp_servers"]]
    no_acl = [d for d in devices if not d["acl"]]
    
    # === ЭТАП 3: ФОРМИРОВАНИЕ ОТЧЁТА ===
    
    report = []
    
    # Заголовок
    report.append(f"""# ОТЧЁТ ПО АУДИТУ ИТ-ИНФРАСТРУКТУРЫ

**Дата:** {datetime.now().strftime('%d.%m.%Y')}
**Основание:** {last_user_prompt[:200]}

---

## 1. ОБЩИЕ СВЕДЕНИЯ

| Параметр | Значение |
|----------|----------|
| Всего устройств | {len(devices)} |
| Ядро/агрегация | {len(core_devices)} |
| Распределение | {len(dist_devices)} |
| Доступ | {len(access_devices)} |
| Периметр/WAN | {len(edge_devices)} |
| Актуальное оборудование | {len(eol_ok)} |
| End-of-Sale | {len(eol_warning)} |
| EOSL (критическое) | {len(eol_critical)} |
| Скриншотов | {len(screenshots_data)} |

""")
    
    # Таблица устройств
    report.append("## 2. СОСТАВ СЕТЕВОГО ОБОРУДОВАНИЯ\n\n")
    report.append("| Hostname | Модель | IP Mgmt | EOL Статус | VLAN | Интерфейсы | AAA | NTP | ACL |\n")
    report.append("|----------|--------|---------|------------|------|------------|-----|-----|-----|\n")
    for d in devices:
        vlan_count = len(d["vlans"])
        iface_count = len(d["interfaces"])
        aaa = "✅" if d["aaa"] else "❌"
        ntp = "✅" if d["ntp_servers"] else "❌"
        acl = "✅" if d["acl"] else "❌"
        eol_icon = "🔴" if d["eol_info"]["status"] == "EOSL" else ("🟡" if d["eol_info"]["status"] == "End-of-Sale" else "🟢")
        report.append(f"| {d['hostname']} | {d['model']} | {d['ip_mgmt']} | {eol_icon} {d['eol_info']['status']} | {vlan_count} | {iface_count} | {aaa} | {ntp} | {acl} |\n")
    
    # Детальный анализ устройств
    report.append("\n## 3. ДЕТАЛЬНЫЙ АНАЛИЗ УСТРОЙСТВ\n\n")
    
    device_details = []
    for d in devices[:15]:
        ospf_str = 'Не настроен'
        if d.get('ospf') and d['ospf'].get('process_id'):
            ospf_str = f"Process {d['ospf'].get('process_id', '?')}, Router ID {d['ospf'].get('router_id', '?')}, Area {d['ospf'].get('area', '?')}"
        
        detail = f"""
### {d['hostname']} ({d['model']})
- **IP управления:** {d['ip_mgmt']}
- **EOL статус:** {d['eol_info']['status']} (срок: {d['eol_info']['eol']})
- **Интерфейсы:** {len(d['interfaces'])} шт.
- **VLAN:** {', '.join([f"{v['id']}({v['ip']})" for v in d['vlans'][:10]]) if d['vlans'] else 'Нет'}
- **Маршруты:** {len(d['routes'])} шт.
- **AAA (TACACS):** {'Настроен' if d['aaa'] else '❌ Не настроен'}{' (серверы: ' + ', '.join(d['tacacs_servers']) + ')' if d['tacacs_servers'] else ''}
- **NTP:** {'Настроен' if d['ntp_servers'] else '❌ Не настроен'}{' (серверы: ' + ', '.join(d['ntp_servers']) + ')' if d['ntp_servers'] else ''}
- **ACL:** {'Настроены' if d['acl'] else '❌ Не настроены'}
- **SNMP:** {'Настроен' if d['snmp'] else '❌ Не настроен'}
- **DHCP Snooping:** {'✅' if d['dhcp_snooping'] else '❌'}
- **Port Security:** {'✅' if d['port_security'] else '❌'}
- **BPDU Protection:** {'✅' if d['bpdu_protection'] else '❌'}
- **OSPF:** {ospf_str}
- **VRRP:** {f"{len(d['vrrp'])} групп" if d['vrrp'] else 'Не настроен'}
- **STP Region:** {d['stp_region'] if d['stp_region'] else 'Не указан'}
- **Syslog:** {', '.join(d['syslog_servers']) if d['syslog_servers'] else 'Не настроен'}
"""
        device_details.append(detail)
    
    device_prompt = f"""Проанализируй следующие сетевые устройства и напиши краткое описание каждого (2-3 предложения) на русском языке:

{chr(10).join(device_details)}

Для каждого устройства укажи его роль в сети и ключевые особенности конфигурации."""
    
    try:
        device_analysis = requests.post(
            f"{OLLAMA_API_URL}/generate",
            json={"model": "qwen2.5-coder:7b", "prompt": device_prompt, "stream": False, "options": {"temperature": 0.3, "num_predict": 3000}},
            timeout=300
        ).json().get("response", "")
        report.append(device_analysis)
    except Exception as e:
        report.append(f"*Анализ устройств не выполнен: {str(e)}*\n")
    
    # EOL раздел
    report.append("\n## 4. АНАЛИЗ ЖИЗНЕННОГО ЦИКЛА ОБОРУДОВАНИЯ\n\n")
    
    if eol_critical:
        report.append("### 4.1 Критическое оборудование (EOSL)\n\n")
        report.append("| Устройство | Модель | EOL Дата | Риск |\n")
        report.append("|------------|--------|----------|------|\n")
        for d in eol_critical:
            report.append(f"| {d['hostname']} | {d['model']} | {d['eol_info']['eol']} | {d['eol_info']['note']} |\n")
    
    if eol_warning:
        report.append("\n### 4.2 Оборудование с ограниченной поддержкой (End-of-Sale)\n\n")
        report.append("| Устройство | Модель | EOL Дата | Рекомендация |\n")
        report.append("|------------|--------|----------|---------------|\n")
        for d in eol_warning:
            report.append(f"| {d['hostname']} | {d['model']} | {d['eol_info']['eol']} | Планировать замену |\n")
    
    # Проблемы безопасности
    report.append("\n## 5. ВЫЯВЛЕННЫЕ ПРОБЛЕМЫ БЕЗОПАСНОСТИ\n\n")
    
    issues = []
    if no_aaa:
        issues.append(f"- **Отсутствует AAA:** {len(no_aaa)} устройств без централизованной аутентификации: {', '.join([d['hostname'] for d in no_aaa])}")
    if no_ntp:
        issues.append(f"- **Отсутствует NTP:** {len(no_ntp)} устройств без синхронизации времени: {', '.join([d['hostname'] for d in no_ntp])}")
    if no_acl:
        issues.append(f"- **Отсутствуют ACL:** {len(no_acl)} устройств без списков контроля доступа: {', '.join([d['hostname'] for d in no_acl])}")
    if eol_critical:
        issues.append(f"- **Критическое устаревание:** {len(eol_critical)} устройств с прекращённой поддержкой (EOSL)")
    
    for issue in issues:
        report.append(issue + "\n")
    
    if not issues:
        report.append("Критических проблем не выявлено.\n")
    
    # Рекомендации
    report.append("\n## 6. РЕКОМЕНДАЦИИ\n\n")
    
    recommendations = []
    if eol_critical:
        recommendations.append("1. **Замена устаревшего оборудования:** Планировать замену устройств EOSL в течение 6-12 месяцев.")
    if no_aaa:
        recommendations.append("2. **Настройка AAA:** Внедрить TACACS+ на всех устройствах для централизованного управления доступом.")
    if no_ntp:
        recommendations.append("3. **Синхронизация времени:** Настроить NTP на всех устройствах, желательно от национального эталона (belgim.by).")
    recommendations.append("4. **Мониторинг:** Настроить централизованный сбор логов (Syslog) и метрик (SNMP) для всех устройств.")
    recommendations.append("5. **Резервное копирование:** Внедрить автоматическое резервное копирование конфигураций.")
    
    for rec in recommendations:
        report.append(rec + "\n")
    
    # Заключение
    report.append("\n## 7. ЗАКЛЮЧЕНИЕ\n\n")
    
    conclusion_prompt = f"""На основе следующего анализа напиши краткое заключение (3-5 предложений) на русском языке:

- Всего устройств: {len(devices)}
- Актуальных: {len(eol_ok)}
- End-of-Sale: {len(eol_warning)}
- EOSL (критическое): {len(eol_critical)}
- Устройств без AAA: {len(no_aaa)}
- Устройств без NTP: {len(no_ntp)}
- Ключевые устройства: {', '.join([d['hostname'] for d in devices[:5]])}
"""
    
    try:
        conclusion = requests.post(
            f"{OLLAMA_API_URL}/generate",
            json={"model": "qwen2.5-coder:7b", "prompt": conclusion_prompt, "stream": False, "options": {"temperature": 0.3, "num_predict": 1000}},
            timeout=120
        ).json().get("response", "")
        report.append(conclusion)
    except:
        report.append("*Заключение не сгенерировано*\n")
    
    # Подвал
    report.append(f"\n\n---\n*Отчёт сгенерирован автоматически {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}*\n")
    report.append(f"*Модель: qwen2.5-coder:7b | Устройств проанализировано: {len(devices)}*\n")
    
    # Сборка и сохранение
    full_report = "\n".join(report)
    
    filename = f"audit_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    path = OUTPUT_DIR / filename
    async with aiofiles.open(path, "w", encoding="utf-8") as f:
        await f.write(full_report)
    
    print(f"REPORT SAVED: {path} ({len(full_report)} chars)")
    print("=" * 60)
    
    return {
        "status": "success",
        "report_file": filename,
        "path": str(path),
        "content_preview": full_report[:500] + "...",
        "devices": len(devices),
        "eol_critical": len(eol_critical),
        "eol_warning": len(eol_warning),
        "issues_found": len(issues)
    }


# ===== ОСТАЛЬНЫЕ ЭНДПОИНТЫ =====

@app.get("/reports")
async def list_reports():
    reports = []
    if OUTPUT_DIR.exists():
        for f in OUTPUT_DIR.iterdir():
            if f.is_file():
                s = f.stat()
                reports.append({"name": f.name, "size_bytes": s.st_size, "modified": datetime.fromtimestamp(s.st_mtime).isoformat()})
    return {"reports": sorted(reports, key=lambda x: x["modified"], reverse=True)}


@app.get("/download/{filename}")
async def download_report(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404)
    return FileResponse(str(path), filename=filename)


@app.get("/view/{filename}")
async def view_report(filename: str):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404)
    async with aiofiles.open(path, "r", encoding="utf-8") as f:
        content = await f.read()
    return {"filename": filename, "content": content}


@app.get("/export/{filename}")
async def export_report(filename: str, format: str = "md"):
    path = OUTPUT_DIR / filename
    if not path.exists():
        raise HTTPException(404)
    async with aiofiles.open(path, "r", encoding="utf-8") as f:
        content = await f.read()
    
    if format == "md":
        return FileResponse(str(path), filename=filename)
    elif format == "txt":
        return PlainTextResponse(content, headers={"Content-Disposition": f"attachment; filename={filename.replace('.md', '.txt')}"})
    elif format == "html":
        html = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        return HTMLResponse(f"<html><head><meta charset='UTF-8'><style>body{{font-family:Arial;max-width:900px;margin:40px auto;padding:20px}}table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #ddd;padding:8px}}</style></head><body>{html}</body></html>")
    else:
        raise HTTPException(400, f"Формат {format} не поддерживается")


@app.get("/reports-page", response_class=HTMLResponse)
async def reports_page(request: Request):
    try:
        template = env.get_template("reports.html")
        return template.render({"request": request, "title": "Просмотр отчётов"})
    except:
        return HTMLResponse("<h2>Отчёты</h2><a href='/'>На главную</a>")
